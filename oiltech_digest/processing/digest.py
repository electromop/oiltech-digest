"""Monthly digest draft generation."""

from __future__ import annotations

import base64
import json
import re
from html import escape
from pathlib import Path
from datetime import UTC, datetime
from urllib.parse import urlparse
from xml.sax.saxutils import escape as xml_escape
from zipfile import ZIP_DEFLATED, ZipFile

from oiltech_digest.config import EXPORTS_DIR
from oiltech_digest.db import repository

TEMPLATE_DIR = Path(__file__).resolve().parent
EMAIL_TEMPLATE = "digest_email_template.html"
BRANDING_CONFIG = "digest_branding.json"
ASSETS_DIR = TEMPLATE_DIR / "assets"
HERO_ASSET = "oiltech_digest_hero_600x360.png"
HERO_ALT = "Нефтесервисный дайджест — технологии, рынок и возможности для бизнеса"

# Корпоративный шрифт GPN Din для PDF (Chromium рендерит @font-face). В email
# оставляем Arial-фолбэк (почтовые клиенты вырезают кастомные шрифты).
_PDF_FONT_FILES = {
    "GPN Din": [("GPN_DIN-Regular.ttf", 400), ("GPN_DIN-Bold.ttf", 700)],
    "GPN Din Condensed": [("GPN_DIN_Condensed-Bold.ttf", 700)],
}


def _asset_bytes(name: str) -> bytes | None:
    path = ASSETS_DIR / name
    return path.read_bytes() if path.exists() else None


def _hero_data_uri() -> str:
    """Утверждённый hero-баннер (600×360) как self-contained data-URI: работает и в
    email, и в PDF (Chromium рендерит из строки, без сервера), и оффлайн."""
    data = _asset_bytes(HERO_ASSET)
    return "data:image/png;base64," + base64.b64encode(data).decode("ascii") if data else ""


def _embedded_font_faces() -> str:
    """@font-face с GPN Din в base64 для HTML/PDF-рендера без зависимости от ОС."""
    faces = []
    for family, files in _PDF_FONT_FILES.items():
        for fname, weight in files:
            data = _asset_bytes(fname)
            if not data:
                continue
            b64 = base64.b64encode(data).decode("ascii")
            faces.append(
                f"@font-face{{font-family:'{family}';font-style:normal;font-weight:{weight};"
                f"src:url(data:font/ttf;base64,{b64}) format('truetype');}}"
            )
    return "".join(faces)


def _pdf_font_face_style() -> str:
    """<style> с GPN Din в base64 — добавляется в PDF-рендер для Chromium."""
    faces = _embedded_font_faces()
    return "<style>" + faces + "</style>" if faces else ""


def _load_digest_branding() -> dict:
    defaults = {
        "header": {
            "brand_text": "ГАЗПРОМ НЕФТЬ",
            "brand_suffix": "ЭНЕРГИЯ В ЛЮДЯХ",
            "department_text": "БЛОК РАЗВИТИЯ БИЗНЕСА",
        },
        "hero": {
            "badge": "НОВОСТИ",
            "headline": "НЕФТЕСЕРВИСНЫЙ ДАЙДЖЕСТ",
            "subtitle": "Технологии, рынок и возможности для бизнеса",
            "image_url": "",
        },
        "issue": {
            "title_template": "Нефтесервисный дайджест",
            "title_template_with_month": "Нефтесервисный дайджест · {month}",
            "period_label_all": "за всё время",
            "preheader": "Ключевые новости и обзоры нефтесервисного рынка",
            "intro_template": (
                "Уважаемые коллеги! Представляем ключевые новости и обзоры нефтесервисного рынка, "
                "которые помогают отслеживать технологические тренды, рыночную динамику "
                "и возможности для развития бизнеса."
            ),
            "intro_template_with_month": (
                "Уважаемые коллеги! Представляем ключевые новости и обзоры за {month}, "
                "которые помогают отслеживать технологические тренды, рыночную динамику "
                "и возможности для развития нефтесервисного бизнеса."
            ),
            "highlights_title": "Главное за период",
            "news_title": "Новости",
            "read_more_label": "ЧИТАТЬ ДАЛЕЕ",
            "empty_summary_text": "Суть ещё не сформирована.",
            "preview_empty_text": "В текущей выборке нет сигналов для превью.",
        },
        "footer": {
            "contact_text": "При возникновении вопросов обращайтесь в Блок развития бизнеса",
            "contact_email": "Rodionov.VVL@gazprom-neft.ru",
            "note": "Внутренняя корпоративная рассылка",
            "socials": [],
        },
        "highlights": {
            "analytics_source_keywords": [
                "mckinsey",
                "wood mac",
                "woodmac",
                "wood mackenzie",
                "rystad",
                "deloitte",
                "iea",
                "eia",
                "bcg",
                "petroleum economist",
                "hart",
                "rigzone",
                "s&p",
                "ihs",
                "mit",
                "bloomberg",
            ],
            "analytics_category_keywords": ["аналит", "обзор", "прогноз", "рынок", "исследов", "report"],
            "business_category_keywords": [
                "бизнес",
                "m&a",
                "сделк",
                "инвест",
                "контракт",
                "возможност",
                "партнёрств",
                "партнерств",
                "локализац",
                "экспорт",
            ],
            "cards": [
                {
                    "metric": "total",
                    "icon": "doc",
                    "prefix": "",
                    "suffix": "",
                    "noun_one": "новость",
                    "noun_few": "новости",
                    "noun_many": "новостей",
                },
                {
                    "metric": "analytics",
                    "icon": "chart",
                    "prefix": "аналитических",
                    "suffix": "",
                    "noun_one": "материал",
                    "noun_few": "материала",
                    "noun_many": "материалов",
                },
                {
                    "metric": "business",
                    "icon": "people",
                    "prefix": "",
                    "suffix": "для бизнеса",
                    "noun_one": "возможность",
                    "noun_few": "возможности",
                    "noun_many": "возможностей",
                },
            ],
        },
    }
    path = TEMPLATE_DIR / BRANDING_CONFIG
    if not path.exists():
        return defaults
    loaded = json.loads(path.read_text(encoding="utf-8"))
    return {
        "header": {**defaults["header"], **(loaded.get("header") or {})},
        "hero": {**defaults["hero"], **(loaded.get("hero") or {})},
        "issue": {**defaults["issue"], **(loaded.get("issue") or {})},
        "footer": {**defaults["footer"], **(loaded.get("footer") or {})},
        "highlights": {**defaults["highlights"], **(loaded.get("highlights") or {})},
    }


def get_digest_branding() -> dict:
    return _load_digest_branding()


def save_digest_branding(payload: dict) -> dict:
    current = _load_digest_branding()
    merged = {
        "header": {**current["header"], **(payload.get("header") or {})},
        "hero": {**current["hero"], **(payload.get("hero") or {})},
        "issue": {**current["issue"], **(payload.get("issue") or {})},
        "footer": {
            **current["footer"],
            **(payload.get("footer") or {}),
            "socials": list((payload.get("footer") or {}).get("socials") or current["footer"].get("socials") or []),
        },
        "highlights": {
            **current["highlights"],
            **(payload.get("highlights") or {}),
        },
    }
    path = TEMPLATE_DIR / BRANDING_CONFIG
    path.write_text(json.dumps(merged, ensure_ascii=False, indent=2), encoding="utf-8")
    return merged


def build_digest_content(
    month: str | None = None,
    limit: int = 20,
    min_score: float = 60,
    user_id: int | None = None,
    max_score: float | None = None,
    search: str | None = None,
    top_tag: str | None = None,
) -> dict:
    branding = _load_digest_branding()
    issue_cfg = branding["issue"]
    rows = []
    if month and user_id is not None:
        saved_digest = repository.get_monthly_digest(month, user_id=user_id)
        saved_ids = [int(item["article_id"]) for item in (saved_digest or {}).get("items", []) if item.get("article_id") is not None]
        if saved_ids:
            rows = repository.digest_items_by_article_ids(saved_ids[:limit], user_id=user_id)
    if not rows:
        rows = repository.digest_candidates(
            month=month,
            limit=limit,
            min_score=min_score,
            user_id=user_id,
            max_score=max_score,
            search=search,
            top_tag=top_tag,
        )
    news = []
    for row in rows:
        tag = row.get("tag_name") or "Без тега"
        if row.get("parent_tag_name"):
            tag = f"{row['parent_tag_name']} / {tag}"
        published = row["published_at"].date().isoformat() if row.get("published_at") else None
        news.append(
            {
                "category": tag,
                "article_id": row["id"],
                "title": row["title"],
                "source": row["source_name"],
                "url": row["url"],
                "published_at": published,
                "tag": tag,
                "score": float(row["total_score"]) if row.get("total_score") is not None else None,
                "score_label": row.get("score_label"),
                "summary": _compact_digest_summary(row.get("summary") or "", row.get("title") or ""),
                "image_url": row.get("image_url") or "",
            }
        )
    title_template = issue_cfg["title_template_with_month"] if month else issue_cfg["title_template"]
    intro_template = issue_cfg["intro_template_with_month"] if month else issue_cfg["intro_template"]
    title = title_template.format(month=month or "")
    intro = intro_template.format(month=month or "")
    return {
        "month": month,
        "title": title,
        "issue": {
            "title": title,
            "period": month or issue_cfg["period_label_all"],
            "preheader": issue_cfg["preheader"],
            "intro": intro,
            "highlights_title": issue_cfg["highlights_title"],
            "news_title": issue_cfg["news_title"],
            "read_more_label": issue_cfg["read_more_label"],
            "empty_summary_text": issue_cfg["empty_summary_text"],
            "preview_empty_text": issue_cfg["preview_empty_text"],
        },
        "hero": {
            "badge": branding["hero"]["badge"],
            "headline": branding["hero"]["headline"],
            "subtitle": branding["hero"]["subtitle"],
            "image_url": branding["hero"]["image_url"],
        },
        "news": news,
        "items": news,
        "highlights": _digest_highlights(news, branding.get("highlights")),
        "footer": {
            "contact_text": branding["footer"]["contact_text"],
            "contact_email": branding["footer"]["contact_email"],
            "note": branding["footer"]["note"],
            "socials": branding["footer"]["socials"],
        },
        "branding": branding,
    }


def render_digest_email(content: dict) -> str:
    """Render the branded Gazprom Neft digest HTML from issue/hero/news/footer.

    The same render is used for the on-screen HTML, the file export and the PDF —
    one template, identical to the reference (digest_email_claude_pack).
    """
    template = (TEMPLATE_DIR / EMAIL_TEMPLATE).read_text(encoding="utf-8")
    news_items = content.get("news", [])
    branding = content.get("branding") or _load_digest_branding()
    news_sections = _render_news_sections(news_items, content.get("issue") or {})
    values = {
        "header_brand_text": _html(branding.get("header", {}).get("brand_text")),
        "header_brand_suffix": _html(branding.get("header", {}).get("brand_suffix")),
        "header_department_text": _html(branding.get("header", {}).get("department_text")),
        "font_face_style": _embedded_font_faces(),
        "issue_title": _html(content.get("issue", {}).get("title")),
        "issue_preheader": _html(content.get("issue", {}).get("preheader")),
        "issue_intro": _html(content.get("issue", {}).get("intro")),
        "highlights_section": "",
        "news_sections": news_sections,
        # Hero — утверждённый баннер: внешний URL из brandinга, иначе встроенная картинка.
        "hero_img_src": _html(content.get("hero", {}).get("image_url")) or _hero_data_uri(),
        "hero_alt": _html(HERO_ALT),
        "footer_contact_text": _html(content.get("footer", {}).get("contact_text")),
        "footer_contact_email": _html(content.get("footer", {}).get("contact_email")),
        "footer_note": _html(content.get("footer", {}).get("note")),
        "footer_socials_html": _render_footer_socials(content.get("footer", {}).get("socials") or branding.get("footer", {}).get("socials") or []),
    }
    return template.format(**values)


# The export is the exact same branded document as the email — one source of truth.
def render_digest_export_html(content: dict) -> str:
    return render_digest_email(content)


def _html(value: object) -> str:
    return escape("" if value is None else str(value), quote=True)


def _compact_digest_summary(summary: str, title: str, max_chars: int = 170) -> str:
    """Shorten AI summary for digest cards so layout stays tight and readable."""
    text = re.sub(r"\s+", " ", (summary or "").strip())
    if not text:
        return ""

    # Our summaries often start with the article title prefix. Remove it for the digest.
    title_prefix = f"{(title or '').strip()}:"
    if title and text.lower().startswith(title_prefix.lower()):
        text = text[len(title_prefix):].strip()

    # Prefer the first sentence if it is already compact enough.
    first_sentence = re.split(r"(?<=[.!?])\s+", text, maxsplit=1)[0].strip()
    candidate = first_sentence or text
    if len(candidate) <= max_chars:
        return candidate

    clipped = candidate[: max_chars - 1].rstrip(" ,;:-")
    if " " in clipped:
        clipped = clipped.rsplit(" ", 1)[0]
    return clipped + "…"


_HL_ICONS = {
    "doc": '<svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="#ffffff" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M7 3h7l4 4v14H7z"/><path d="M14 3v4h4"/><line x1="10" y1="13" x2="15" y2="13"/><line x1="10" y1="17" x2="15" y2="17"/></svg>',
    "chart": '<svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="#ffffff" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polyline points="4 16 9 11 13 14 20 6"/><polyline points="15 6 20 6 20 11"/></svg>',
    "people": '<svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="#ffffff" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="9" cy="9" r="3"/><path d="M3.5 20a5.5 5.5 0 0 1 11 0"/><path d="M16 7.5a3 3 0 0 1 0 5.5"/><path d="M15.5 14.5a5.5 5.5 0 0 1 5 5.5"/></svg>',
}


def _plural(n: int, one: str, few: str, many: str) -> str:
    """Русское склонение существительного по числу (1 новость / 2 новости / 5 новостей)."""
    n = abs(int(n)) % 100
    if 10 < n < 20:
        return many
    n1 = n % 10
    if n1 == 1:
        return one
    if 2 <= n1 <= 4:
        return few
    return many


def _digest_highlights(news: list[dict], config: dict | None = None) -> list[dict]:
    """KPI выпуска для блока «Главное за период»: всего новостей, аналитических
    материалов, бизнес-возможностей. Аналитику/возможности оцениваем эвристикой
    по источнику/категории (правила можно уточнить позже)."""
    defaults = _load_digest_branding().get("highlights") or {}
    config = {**defaults, **(config or {})}

    def cat(n: dict) -> str:
        return (n.get("category") or "").lower()

    def src(n: dict) -> str:
        return (n.get("source") or "").lower()

    analytic_src = tuple(str(item).lower() for item in config.get("analytics_source_keywords") or [])
    analytic_kw = tuple(str(item).lower() for item in config.get("analytics_category_keywords") or [])
    business_kw = tuple(str(item).lower() for item in config.get("business_category_keywords") or [])
    analytics = sum(
        1
        for n in news
        if any(k in cat(n) for k in analytic_kw) or any(s in src(n) for s in analytic_src)
    )
    business = sum(1 for n in news if any(k in cat(n) for k in business_kw))
    total = len(news)
    metric_values = {
        "total": total,
        "analytics": analytics,
        "business": business,
    }
    cards = config.get("cards") or []
    if not cards:
        cards = _load_digest_branding()["highlights"]["cards"]
    result = []
    for card in cards:
        metric = str(card.get("metric") or "total")
        value = int(metric_values.get(metric, 0))
        prefix = str(card.get("prefix") or "").strip()
        suffix = str(card.get("suffix") or "").strip()
        noun = _plural(
            value,
            str(card.get("noun_one") or "материал"),
            str(card.get("noun_few") or "материала"),
            str(card.get("noun_many") or "материалов"),
        )
        label = " ".join(part for part in [prefix, noun, suffix] if part).strip()
        result.append(
            {
                "metric": metric,
                "value": value,
                "icon": str(card.get("icon") or "doc"),
                "label": label,
            }
        )
    return result


def _render_highlights(highlights: list[dict]) -> str:
    """KPI-плашки: иконка + крупное число + подпись."""
    if not highlights:
        return ""
    width = f"{max(1, int(100 / max(1, len(highlights))))}%"
    cells = []
    for h in highlights:
        icon = _HL_ICONS.get(h.get("icon", ""), "")
        cells.append(
            f'<td width="{width}" valign="top" style="padding:0 10px 0 0;">'
            '<table role="presentation" cellspacing="0" cellpadding="0" border="0"><tr>'
            '<td valign="middle" width="58">'
            '<div style="width:46px;height:46px;background:#003da6;border-radius:8px;text-align:center;line-height:46px;">'
            f'{icon}</div></td>'
            '<td valign="middle" style="padding-left:10px;">'
            f'<div style="font-size:26px;line-height:28px;color:#003da6;font-weight:bold;">{_html(h.get("value"))}</div>'
            f'<div class="highlights-label">{_html(h.get("label"))}</div>'
            '</td></tr></table></td>'
        )
    return (
        '<table role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0">'
        f'<tr>{"".join(cells)}</tr></table>'
    )


def _render_highlights_section(issue: dict, highlights: list[dict]) -> str:
    if not highlights:
        return ""
    title = issue.get("highlights_title") or "Главное за период"
    highlights_html = _render_highlights(highlights)
    return f"""
          <tr>
            <td style="padding:8px 36px 14px 36px;">
              <table class="highlights-card" role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0" style="border:1px solid #d9e3f3;border-radius:8px;background:#f7faff;">
                <tr>
                  <td style="padding:18px 18px 16px 18px;">
                    <div style="font-family:'GPN Din Condensed','GPN Din',Arial,Helvetica,sans-serif;font-size:30px;line-height:36px;color:#003da6;font-weight:bold;letter-spacing:.10em;text-transform:uppercase;margin-bottom:16px;">{_html(title)}</div>
                    {highlights_html}
                  </td>
                </tr>
              </table>
            </td>
          </tr>"""


def _render_news_sections(news_items: list[dict], issue: dict) -> str:
    title = issue.get("news_title") or "Новости"
    if not news_items:
        return f'<div class="news-section-title">{_html(title)}</div>'
    items_html = "\n".join(
        '<tr><td style="padding:0;">'
        f'{_render_news_item(item, issue)}'
        "</td></tr>"
        for item in news_items
    )
    return (
        '<div class="news-page">'
        '<table class="news-repeat-table" role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0">'
        f'<thead><tr><td style="padding:0;"><div class="news-section-title">{_html(title)}</div></td></tr></thead>'
        f'<tbody>{items_html}</tbody>'
        "</table>"
        "</div>"
    )


def _render_news_item(item: dict, issue: dict | None = None) -> str:
    """Карточка новости (формат референса коллеги): фото + заголовок сверху,
    ниже — краткое описание и строка «Читать далее | теги». Карточка не должна
    разрываться между страницами (CSS .news-card { break-inside: avoid })."""
    issue = issue or {}
    image_url = item.get("image_url") or ""
    summary = item.get("summary") or issue.get("empty_summary_text") or "Суть ещё не сформирована."
    read_more_label = issue.get("read_more_label") or "ЧИТАТЬ ДАЛЕЕ"
    meta = _render_news_meta(item)
    if image_url and not _is_unusable_digest_image_url(image_url):
        media = (
            f'<img class="news-card-image" src="{_html(image_url)}" width="130" height="86" alt="{_html(item.get("title"))}" '
            'style="display:block;width:130px;height:86px;object-fit:cover;border-radius:6px;border:0;">'
        )
    else:
        # Для статей без изображения рисуем стабильную заглушку по верхнему тегу:
        # карточки одной рубрики всегда получают одинаковый фон/подпись.
        placeholder = _news_placeholder_data_uri(item.get("category"))
        media = (
            f'<img class="news-card-image" src="{placeholder}" width="130" height="86" alt="" '
            'style="display:block;width:130px;height:86px;border-radius:6px;border:0;">'
        )
    return f"""
              <table class="news-card" role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0" style="margin:0 0 18px 0;border:1px solid #d9e3f3;border-radius:8px;background:#ffffff;">
                <tr>
                  <td width="158" valign="top" style="padding:12px 14px 8px 12px;">
                    {media}
                  </td>
                  <td valign="top" style="padding:14px 16px 8px 0;">
                    <div class="news-card-title">{_html(item.get("title"))}</div>
                  </td>
                </tr>
                <tr>
                  <td colspan="2" valign="top" style="padding:0 16px 14px 16px;">
                    <div class="news-card-summary">{_html(summary)}</div>
                    {meta}
                    <table role="presentation" width="100%" cellspacing="0" cellpadding="0" border="0" style="margin-top:8px;">
                      <tr>
                        <td align="left" valign="middle">
                          <a href="{_html(item.get("url"))}" style="color:#e83d08;text-decoration:none;font-size:13px;line-height:18px;font-weight:bold;letter-spacing:.06em;text-transform:uppercase;">{_html(read_more_label)} &#8594;</a>
                        </td>
                        <td align="right" valign="middle">
                          <span class="news-card-tag">{_html(item.get("category"))}</span>
                        </td>
                      </tr>
                    </table>
                  </td>
                </tr>
              </table>"""


def _render_news_meta(item: dict) -> str:
    parts = []
    source = str(item.get("source") or "").strip()
    if source:
        parts.append(source)
    published = _format_digest_date(item.get("published_at"))
    if published:
        parts.append(published)
    if not parts:
        return ""
    return f'<div class="news-card-meta">{_html(" · ".join(parts))}</div>'


def _render_footer_socials(socials: list[dict]) -> str:
    if not socials:
        return ""
    cells = []
    for item in socials:
        icon = _footer_social_icon_svg(item)
        cells.append(
            '<td style="padding:0 7px;">'
            '<div style="width:36px;height:36px;background:#ffffff;border-radius:50%;'
            'text-align:center;line-height:36px;display:inline-block;'
            f'color:{_html(item.get("accent") or "#262d3c")};'
            "font-weight:bold;font-size:13px;font-family:'GPN Din',Arial,Helvetica,sans-serif;"
            f'" title="{_html(item.get("label"))}">{icon}</div>'
            "</td>"
        )
    return "".join(cells)


def _footer_social_icon_svg(item: dict) -> str:
    label = str(item.get("label") or "").strip().lower()
    text = str(item.get("text") or "").strip()
    accent = _html(item.get("accent") or "#262d3c")
    key = re.sub(r"[^a-zа-я0-9+]+", "", label, flags=re.IGNORECASE)
    if key in {"vk", "вк"}:
        return (
            '<svg width="36" height="36" viewBox="0 0 36 36" xmlns="http://www.w3.org/2000/svg" '
            'style="display:block;">'
            f'<path fill="{accent}" d="M9.8 12.3h3.1c.1 3.1 1.4 5 2.5 5.4v-5.4h2.9v3.1c1.1-.1 2.2-1.7 2.6-3.1h3.1c-.5 1.9-1.8 3.4-3 4.1 1.3.6 2.7 2 3.3 4.4h-3.4c-.4-1.4-1.4-2.5-2.6-2.7v2.7h-.4c-5.2 0-7.9-3.4-8.1-8.5z"/>'
            "</svg>"
        )
    if key in {"telegram", "tg", "телеграм"}:
        return (
            '<svg width="36" height="36" viewBox="0 0 36 36" xmlns="http://www.w3.org/2000/svg" '
            'style="display:block;">'
            f'<path fill="{accent}" d="M25.9 10.7 8.7 17.3c-1.2.5-1.2 1.1-.2 1.4l4.4 1.4 1.7 5.2c.2.6.3.8.7.8.4 0 .6-.2.9-.5l2.1-2 4.4 3.2c.8.4 1.3.2 1.5-.7l2.8-13.3c.3-1.1-.4-1.6-1.1-1.3zm-3.8 3.1-7.9 7.1-.3 3.1-1.2-4.1 9.4-6.1z"/>'
            "</svg>"
        )
    if key in {"youtube", "ютуб"}:
        return (
            '<svg width="36" height="36" viewBox="0 0 36 36" xmlns="http://www.w3.org/2000/svg" '
            'style="display:block;">'
            f'<path fill="{accent}" d="M26.1 13.5c-.2-.9-.9-1.6-1.8-1.8C22.8 11.3 18 11.3 18 11.3s-4.8 0-6.3.4c-.9.2-1.6.9-1.8 1.8-.4 1.6-.4 4.9-.4 4.9s0 3.3.4 4.9c.2.9.9 1.6 1.8 1.8 1.5.4 6.3.4 6.3.4s4.8 0 6.3-.4c.9-.2 1.6-.9 1.8-1.8.4-1.6.4-4.9.4-4.9s0-3.3-.4-4.9z"/>'
            '<path fill="#fff" d="m16.4 21.3 5-2.9-5-2.9v5.8z"/>'
            "</svg>"
        )
    if key in {"rt", "rutube", "рутуб"}:
        return (
            '<svg width="36" height="36" viewBox="0 0 36 36" xmlns="http://www.w3.org/2000/svg" '
            'style="display:block;">'
            f'<rect x="8" y="12" width="20" height="12" rx="3" fill="{accent}"/>'
            '<path fill="#fff" d="M15.2 20.8v-5.4h4.2c1.5 0 2.4.8 2.4 2 0 1-.6 1.6-1.4 1.9l1.5 1.5h-2.3l-1.2-1.3h-1.4v1.3h-1.8zm1.8-2.8h2.2c.5 0 .8-.2.8-.6s-.3-.6-.8-.6H17v1.2z"/>'
            "</svg>"
        )
    if key in {"дзен", "dzen", "zen"}:
        return (
            '<svg width="36" height="36" viewBox="0 0 36 36" xmlns="http://www.w3.org/2000/svg" '
            'style="display:block;">'
            f'<circle cx="18" cy="18" r="9" fill="{accent}"/>'
            '<path fill="#fff" d="M17 9.5h2c0 4.1 2.4 6.5 6.5 6.5v2c-4.1 0-6.5 2.4-6.5 6.5h-2c0-4.1-2.4-6.5-6.5-6.5v-2c4.1 0 6.5-2.4 6.5-6.5z"/>'
            "</svg>"
        )
    fallback = _html(text or item.get("label") or "")
    return (
        '<svg width="36" height="36" viewBox="0 0 36 36" xmlns="http://www.w3.org/2000/svg" '
        'style="display:block;">'
        f'<text x="18" y="21.5" text-anchor="middle" font-family="Arial,Helvetica,sans-serif" '
        f'font-size="11" font-weight="700" fill="{accent}">{fallback}</text>'
        "</svg>"
    )


def _news_placeholder_data_uri(category: object) -> str:
    raw = str(category or "Новости").strip()
    top = raw.split(" / ", 1)[0].strip() or "Новости"
    key = re.sub(r"[^a-zа-я0-9]+", "", top.lower(), flags=re.IGNORECASE)
    palette = {
        "технологии": ("#003DA6", "#0057D9"),
        "рынок": ("#0A5C36", "#0F8A50"),
        "бизнессигналы": ("#7A2E0B", "#E83D08"),
        "россия": ("#3F3A8C", "#625CDA"),
        "международноесотрудничество": ("#005B66", "#0097A7"),
        "бурение": ("#004A99", "#1D74D1"),
    }
    start, end = palette.get(key, ("#003DA6", "#001D50"))
    badge_lines = _placeholder_badge_lines(top)
    badge_width = min(114, max(42, max(len(line) for line in badge_lines) * 5.5 + 16))
    badge_height = 35 if len(badge_lines) > 2 else 26 if len(badge_lines) > 1 else 17
    badge_text = "".join(
        f"<text x='14' y='{21 + idx * 9}' font-family='Arial,Helvetica,sans-serif' font-size='7.5' font-weight='700' fill='#ffffff'>{escape(line, quote=True)}</text>"
        for idx, line in enumerate(badge_lines)
    )
    svg = (
        "<svg xmlns='http://www.w3.org/2000/svg' width='130' height='86' viewBox='0 0 130 86'>"
        "<defs><linearGradient id='g' x1='0' y1='0' x2='1' y2='1'>"
        f"<stop offset='0' stop-color='{start}'/><stop offset='1' stop-color='{end}'/>"
        "</linearGradient></defs>"
        "<rect width='130' height='86' rx='6' fill='url(#g)'/>"
        f"<rect x='10' y='11' width='{badge_width}' height='{badge_height}' rx='8.5' fill='rgba(255,255,255,0.18)'/>"
        f"{badge_text}"
        "</svg>"
    )
    return "data:image/svg+xml;base64," + base64.b64encode(svg.encode("utf-8")).decode("ascii")


def _placeholder_badge_lines(label: str, *, max_chars: int = 15, max_lines: int = 3) -> list[str]:
    """Short readable placeholder label without cutting words in the middle."""
    words = re.findall(r"[A-Za-zА-Яа-яЁё0-9]+", str(label).upper())
    words = [word for word in words if word not in {"И", "AND"}]
    if not words:
        return ["НОВОСТИ"]
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if current and len(candidate) > max_chars:
            lines.append(current)
            current = word
            if len(lines) == max_lines:
                break
        else:
            current = candidate
        if len(lines) == max_lines:
            break
    if current and len(lines) < max_lines:
        lines.append(current)
    return lines[:max_lines] or ["НОВОСТИ"]


def _is_unusable_digest_image_url(url: object) -> bool:
    """Явно тестовые/локальные/пустые адреса не рендерим как реальные картинки."""
    if not isinstance(url, str) or not url.strip():
        return True
    raw = url.strip()
    parsed = urlparse(raw)
    if parsed.scheme not in {"http", "https", "data"}:
        return True
    if parsed.scheme == "data":
        return False
    host = (parsed.hostname or "").lower()
    if host in {"example.com", "www.example.com", "localhost", "127.0.0.1"}:
        return True
    return False


def _format_digest_date(value: object) -> str:
    if not value:
        return ""
    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return ""
        try:
            if len(raw) >= 10:
                dt = datetime.fromisoformat(raw[:10])
                return dt.strftime("%d.%m.%Y")
        except ValueError:
            return raw
        return raw
    if hasattr(value, "strftime"):
        try:
            return value.strftime("%d.%m.%Y")
        except Exception:
            return str(value)
    return str(value)


def write_digest_content(path: str | Path, month: str, limit: int = 20,
                         min_score: float = 60, html_path: str | Path | None = None,
                         user_id: int | None = None, max_score: float | None = None,
                         search: str | None = None, top_tag: str | None = None) -> dict:
    content = build_digest_content(
        month=month,
        limit=limit,
        min_score=min_score,
        user_id=user_id,
        max_score=max_score,
        search=search,
        top_tag=top_tag,
    )
    output_path = Path(path)
    output_path.write_text(json.dumps(content, ensure_ascii=False, indent=2), encoding="utf-8")
    result = {"path": str(output_path), "items": len(content["items"])}
    if html_path:
        html_output_path = Path(html_path)
        html_output_path.write_text(render_digest_email(content), encoding="utf-8")
        result["html_path"] = str(html_output_path)
    return result


def save_digest_draft(
    month: str,
    limit: int = 20,
    min_score: float = 60,
    user_id: int | None = None,
    max_score: float | None = None,
    search: str | None = None,
    top_tag: str | None = None,
) -> dict:
    """Build digest content from current selected articles and persist it as a draft."""
    content = build_digest_content(
        month=month,
        limit=limit,
        min_score=min_score,
        user_id=user_id,
        max_score=max_score,
        search=search,
        top_tag=top_tag,
    )
    items = [
        {
            "article_id": item["article_id"],
            "section": item.get("category"),
            "editor_note": item.get("summary"),
        }
        for item in content.get("items", [])
        if item.get("article_id") is not None
    ]
    saved = repository.save_monthly_digest(
        month=month,
        title=content.get("title") or f"Нефтесервисный дайджест · {month}",
        items=items,
        status="draft",
        user_id=user_id,
    )
    return {**saved, "content_items": len(content.get("items", []))}


def render_digest_pdf(content: dict) -> bytes:
    """Render the branded digest to PDF via headless Chromium (pixel-perfect).

    Playwright/Chromium is an optional, server-side dependency (it is heavy and not
    needed for tests or for the HTML/Word paths), so it is imported lazily and a
    clear, actionable error is raised when it is missing.
    """
    html_str = render_digest_email(content)
    # Вшиваем GPN Din только в PDF: Chromium рендерит @font-face, корпоративный шрифт
    # попадает в PDF (в email он не нужен — клиенты вырезают кастомные шрифты).
    font_style = _pdf_font_face_style()
    if font_style:
        html_str = html_str.replace("</head>", font_style + "</head>", 1)
    try:
        from playwright.sync_api import sync_playwright
    except ImportError as exc:  # pragma: no cover - depends on optional dep
        raise RuntimeError(
            "PDF-экспорт требует Playwright с Chromium. Установите на сервере: "
            "pip install playwright && python -m playwright install --with-deps chromium"
        ) from exc

    with sync_playwright() as pw:
        browser = pw.chromium.launch(args=["--no-sandbox"])
        try:
            page = browser.new_page()
            page.set_content(html_str, wait_until="load")
            pdf_bytes = page.pdf(
                format="A4",
                print_background=True,
                margin={"top": "0", "bottom": "0", "left": "0", "right": "0"},
            )
        finally:
            browser.close()
    return pdf_bytes


def _fetch_docx_image(url: str | None) -> bytes | None:
    """Скачать картинку статьи для вставки в Word. Любая ошибка/неподходящий тип → None (пропуск)."""
    if not url or not isinstance(url, str) or not url.startswith(("http://", "https://")):
        return None
    try:
        import requests

        resp = requests.get(url, timeout=8, headers={"User-Agent": "Mozilla/5.0 OilTechDigest"})
        content_type = resp.headers.get("content-type", "")
        if resp.ok and content_type.startswith("image/") and "svg" not in content_type:
            data = resp.content
            if 0 < len(data) <= 8_000_000:  # не тащим гигантские файлы в документ
                return data
    except Exception:
        return None
    return None


def _docx_hero_bytes(hero: dict | None = None) -> bytes | None:
    hero = hero or {}
    image_url = hero.get("image_url")
    data = _fetch_docx_image(image_url) if image_url else None
    if data:
        return data
    asset = _asset_bytes(HERO_ASSET)
    return asset if asset else None


def _add_docx_hyperlink(paragraph, url: str, text: str, color_hex: str | None = None) -> None:
    """Вставить кликабельную ссылку в параграф python-docx (нативной поддержки нет)."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    if color_hex:
        color = OxmlElement("w:color")
        color.set(qn("w:val"), color_hex)
        rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _chunk_news_items(news_items: list[dict], size: int = 3) -> list[list[dict]]:
    """Разбить новости на «страницы» по size карточек — для постраничного DOCX
    (render_digest_docx делит новости по 3 с разрывом страницы между чанками).
    HTML-путь больше не использует этот хелпер, но DOCX-путь зависит от него —
    не удалять без правки render_digest_docx."""
    if size <= 0:
        size = 3
    return [news_items[index:index + size] for index in range(0, len(news_items), size)]


def render_digest_docx(content: dict) -> bytes:
    """Render the branded digest to a rich .docx via python-docx.

    Включает шапку бренда, hero, новости с фото, кликабельные ссылки
    «Читать далее» и футер.
    """
    from io import BytesIO

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    issue = content.get("issue") or {}
    hero = content.get("hero") or {}
    news_items = content.get("news") or content.get("items") or []
    footer = content.get("footer") or {}
    header = (content.get("branding") or {}).get("header") or {}

    BRAND_HEX = "003DA6"
    ACCENT_HEX = "E83D08"
    BRAND = RGBColor(0x00, 0x3D, 0xA6)
    ACCENT = RGBColor(0xE8, 0x3D, 0x08)
    GREY = RGBColor(0x66, 0x66, 0x66)

    doc = Document()
    # Корпоративный шрифт во всём документе (рендерится у получателей с установленным
    # GPN Din; иначе Word подставит замену). Базовый стиль Normal наследуют все абзацы.
    doc.styles["Normal"].font.name = "GPN Din"
    if "Title" in doc.styles:
        doc.styles["Title"].font.name = "GPN Din Condensed"
        doc.styles["Title"].font.bold = True
    if "Heading 1" in doc.styles:
        doc.styles["Heading 1"].font.name = "GPN Din Condensed"
        doc.styles["Heading 1"].font.bold = True
    if "Heading 2" in doc.styles:
        doc.styles["Heading 2"].font.name = "GPN Din Condensed"
        doc.styles["Heading 2"].font.bold = True
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.6)
    section.left_margin = section.right_margin = Inches(0.7)

    # --- Шапка бренда ---
    if header.get("brand_text"):
        p = doc.add_paragraph()
        run = p.add_run(header.get("brand_text", ""))
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = BRAND
        if header.get("brand_suffix"):
            suffix = p.add_run("   " + header["brand_suffix"])
            suffix.font.size = Pt(9)
            suffix.font.color.rgb = GREY
    if header.get("department_text"):
        p = doc.add_paragraph()
        run = p.add_run(header["department_text"])
        run.font.size = Pt(9)
        run.font.color.rgb = GREY

    # --- Hero ---
    if hero.get("badge"):
        p = doc.add_paragraph()
        run = p.add_run(hero["badge"])
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = ACCENT
    hero_bytes = _docx_hero_bytes(hero)
    if hero_bytes:
        try:
            doc.add_picture(BytesIO(hero_bytes), width=Inches(6.0))
        except Exception:
            pass
    doc.add_heading(hero.get("headline") or issue.get("title") or "Нефтесервисный дайджест", level=0)
    if hero.get("subtitle"):
        p = doc.add_paragraph()
        run = p.add_run(hero["subtitle"])
        run.italic = True
        run.font.size = Pt(11)
    if issue.get("intro"):
        doc.add_paragraph(issue["intro"])

    # --- Новости ---
    read_more = issue.get("read_more_label") or "Читать далее"
    news_chunks = _chunk_news_items(news_items, size=3)
    for chunk_index, chunk in enumerate(news_chunks):
        if chunk_index:
            doc.add_page_break()
        doc.add_heading(issue.get("news_title") or "Новости", level=1)
        for index, item in enumerate(chunk, start=1 + chunk_index * 3):
            doc.add_heading(item.get("title") or f"Материал {index}", level=2)
            image = _fetch_docx_image(item.get("image_url"))
            if image:
                try:
                    doc.add_picture(BytesIO(image), width=Inches(2.8))
                except Exception:
                    pass
            meta_parts = []
            if item.get("category"):
                meta_parts.append(str(item["category"]))
            if item.get("source"):
                meta_parts.append(str(item["source"]))
            published = _format_digest_date(item.get("published_at"))
            if published:
                meta_parts.append(published)
            if meta_parts:
                p = doc.add_paragraph()
                run = p.add_run(" · ".join(meta_parts))
                run.font.size = Pt(9)
                run.font.color.rgb = GREY
            if item.get("summary"):
                doc.add_paragraph(item["summary"])
            if item.get("url"):
                p = doc.add_paragraph()
                _add_docx_hyperlink(p, str(item["url"]), f"{read_more} →", ACCENT_HEX)

    # --- Футер ---
    doc.add_paragraph()
    if footer.get("contact_text") or footer.get("contact_email"):
        p = doc.add_paragraph()
        if footer.get("contact_text"):
            p.add_run(footer["contact_text"] + " ")
        if footer.get("contact_email"):
            _add_docx_hyperlink(p, f"mailto:{footer['contact_email']}", str(footer["contact_email"]), BRAND_HEX)
    socials = footer.get("socials") or []
    labels = " · ".join(s.get("label", "") for s in socials if s.get("label"))
    if labels:
        p = doc.add_paragraph()
        run = p.add_run(labels)
        run.font.size = Pt(9)
        run.font.color.rgb = GREY
    if footer.get("note"):
        p = doc.add_paragraph()
        run = p.add_run(footer["note"])
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def write_digest_export(
    month: str | None = None,
    export_format: str = "pdf",
    limit: int = 100,
    min_score: float = 0,
    user_id: int | None = None,
    max_score: float | None = None,
    search: str | None = None,
    top_tag: str | None = None,
) -> dict:
    content = build_digest_content(
        month=month,
        limit=limit,
        min_score=min_score,
        user_id=user_id,
        max_score=max_score,
        search=search,
        top_tag=top_tag,
    )
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    base_name = f"digest-{month or 'all'}-{stamp}"
    normalized_format = "docx" if export_format == "doc" else export_format

    if normalized_format == "pdf":
        path = EXPORTS_DIR / f"{base_name}.pdf"
        path.write_bytes(render_digest_pdf(content))
        media_type = "application/pdf"
    elif normalized_format == "docx":
        path = EXPORTS_DIR / f"{base_name}.docx"
        path.write_bytes(render_digest_docx(content))
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif normalized_format == "json":
        path = EXPORTS_DIR / f"{base_name}.json"
        path.write_text(json.dumps(content, ensure_ascii=False, indent=2), encoding="utf-8")
        media_type = "application/json"
    else:  # html
        path = EXPORTS_DIR / f"{base_name}.html"
        path.write_text(render_digest_export_html(content), encoding="utf-8")
        media_type = "text/html; charset=utf-8"

    return {
        "path": str(path),
        "filename": path.name,
        "media_type": media_type,
        "items": len(content["items"]),
        "format": normalized_format,
    }
